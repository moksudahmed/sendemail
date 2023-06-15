import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
import re
from docx import Document

def replace_text_in_word_file(file_path, old_text, new_text):
    # Load the Word document
    doc = Document(file_path)

    # Replace the old text with the new text
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    # Save the modified document
    doc.save(file_path)


def replace_job_bank_number(file_path, new_job_bank_number):
    # Load the Word document
    doc = Document(file_path)

    # Define the pattern to find the job bank number
    pattern = r'Job Bank\s*:\s*\d+'

    # Iterate through the paragraphs and replace the job bank number
    for paragraph in doc.paragraphs:
        for match in re.finditer(pattern, paragraph.text):
            old_job_bank_number = match.group(0)
            paragraph.text = paragraph.text.replace(old_job_bank_number, f'Job Bank : {new_job_bank_number}')

    # Save the modified document
    doc.save(file_path)

def prepare_email_attachment():
    #change()
    file_path = 'Cover Letter.docx'  # Replace with the path to your Word file
    file_path2 = 'cv@moksud.docx'  # Replace with the path to your Word file
    old_text = input("Enter Old Company Name : ") #'Metropolitan'  # Text to be replaced
    new_text = input("Enter New Company Name : ") #'Abdulnour-Escobedo United Inc'  # New text
    new_job_bank_number = input("Enter New Job Number : ") #'Abdulnour-Escobedo United Inc'  # New text
    replace_text_in_word_file(file_path, old_text, new_text)    
    replace_job_bank_number(file_path, new_job_bank_number)
    print("Finished!")



def send_email_with_attachments(sender_email, sender_password, receiver_email, subject, message, attachment_paths):
    # Create the email message
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = receiver_email
    msg['Subject'] = subject

    # Add the message to the email
    msg.attach(MIMEText(message, 'plain'))

    # Attach the attachments to the email
    for attachment_path in attachment_paths:
        with open(attachment_path, "rb") as attachment:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename= {attachment_path}')
        msg.attach(part)

    # Establish a secure connection with the SMTP server
    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()

    # Login to the sender's email account
    server.login(sender_email, sender_password)

    # Send the email
    server.sendmail(sender_email, receiver_email, msg.as_string())

    # Close the connection to the SMTP server
    server.quit()



if __name__ == '__main__':   
    # Example usage
    sender_email = ""
    sender_password = ""
    receiver_email = ""
    subject = "Application for job"
    message = "Hello, please find the attached file."
    attachment_path = "Cover_Letter.docx"  # Replace with the path to your DOCX file
    attachment_paths = ["", ""]  # Replace with the paths to your DOCX files

    send_email_with_attachments(sender_email, sender_password, receiver_email, subject, message, attachment_paths)
