import re
from docx import Document

def replace_text_in_word_file(file_path, old_text, new_text):
    # Load the Word document
    doc = Document(file_path)

    # Replace the old text with the new text
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    # Save the modified document
    doc.save(file_path)


def replace_job_bank_number(file_path, new_job_bank_number):
    # Load the Word document
    doc = Document(file_path)

    # Define the pattern to find the job bank number
    pattern = r'Job Bank\s*:\s*\d+'

    # Iterate through the paragraphs and replace the job bank number
    for paragraph in doc.paragraphs:
        for match in re.finditer(pattern, paragraph.text):
            old_job_bank_number = match.group(0)
            paragraph.text = paragraph.text.replace(old_job_bank_number, f'Job Bank : {new_job_bank_number}')

    # Save the modified document
    doc.save(file_path)

if __name__ == '__main__':   

    #change()
    file_path = 'Cover Letter.docx'  # Replace with the path to your Word file
    file_path2 = 'cv@moksud.docx'  # Replace with the path to your Word file
    old_text = input("Enter Old Company Name : ") #'Metropolitan'  # Text to be replaced
    new_text = input("Enter New Company Name : ") #'Abdulnour-Escobedo United Inc'  # New text
    new_job_bank_number = input("Enter New Job Number : ") #'Abdulnour-Escobedo United Inc'  # New text

    replace_text_in_word_file(file_path, old_text, new_text)    
    replace_job_bank_number(file_path, new_job_bank_number)


    print("Finished!")


